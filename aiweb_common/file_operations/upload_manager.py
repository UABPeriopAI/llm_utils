import base64
import os
import tempfile
from abc import abstractmethod
from io import BytesIO
from pathlib import Path
from typing import Any, Union
import pandas as pd
import pypandoc
import streamlit as st
from docx import Document
from fastapi import BackgroundTasks, HTTPException
from aiweb_common.file_operations.file_handling import ingest_docx_bytes

class UploadManager:
    @abstractmethod
    def read_file(self, file, extension):
        raise NotImplementedError
    @abstractmethod
    def upload_file(self):
        raise NotImplementedError

    def read_pdf(self, file, document_analysis_client):
        poller = document_analysis_client.begin_analyze_document(
            model_id="prebuilt-read", document=file
        )
        result = poller.result()
        text = ""
        for page in result.pages:
            for line in page.lines:
                text += line.content + "\n"
        return text

class StreamlitUploadManager(UploadManager):
    def __init__(
        self,
        file=None,
        message: str = "Please upload a file",
        file_types: list = None,
        accept_multiple_files: bool = False,
        document_analysis_client=None,
    ):
        """
        Allows either an already-uploaded file (passed via `file`) or performs an interactive upload.

        Args:
            file: (Optional) an already-uploaded file object.
            message: The label for the uploader widget.
            file_types: List of allowed file extensions (default list if None).
            accept_multiple_files: Whether to allow multiple file uploads.
            document_analysis_client: (Optional) any additional client if needed.
        """
        self.file = file
        self.message = message
        self.file_types = file_types if file_types is not None else ["csv", "xlsx", "docx", "pdf", "txt"]
        self.accept_multiple_files = accept_multiple_files
        self.document_analysis_client = document_analysis_client

    def process_upload(self):
        """
        If no file has been provided during initialization, show the file uploader.
        Then, process the file based on its extension.
        Returns a tuple (processed_file, extension) or (None, None) if no file is provided.
        """
        # If no file was provided externally then invoke the uploader.
        if self.file is None:
            self.file = st.file_uploader(
                label=self.message, 
                type=self.file_types, 
                accept_multiple_files=self.accept_multiple_files
            )
            if not self.file:
                st.write("Please upload a file to continue...")
                return None, None
        # Call read_file on the provided file.
        extension = Path(self.file.name).suffix
        return self.read_file(self.file, extension=extension)

    def upload_file(self):
        """
        Wraps process_upload for backward compatibility. You can choose your naming.
        """
        return self.process_upload()

    def read_file(self, file, extension):
        # print("Extension - ", extension)
        if extension == ".xlsx":
            return pd.read_excel(file), extension
        elif extension == ".docx":
            doc = Document(file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text, extension
        elif extension == ".csv":
            return pd.read_csv(file), extension
        elif extension == ".pdf":
            return self.read_pdf(file, self.document_analysis_client), extension
        else:
            return None, None


class FastAPIUploadManager(UploadManager):
    def __init__(self, background_tasks: BackgroundTasks):
        self.background_tasks = background_tasks

    def process_file_bytes(
        self, file: bytes, extension: str
    ) -> Union[pd.DataFrame, str]:
        """
        Reads the file from byte string based on the file extension and returns
        either a DataFrame or a markdown string.

        Args:
            file_bytes (bytes): The byte-encoded content of the file.
            extension (str): The file extension indicating the file type.

        Returns:
            Union[pd.DataFrame, str]: Depending on the file extension, either returns a DataFrame for Excel files
            or a markdown string for DOCX and TXT files.
        """
        print("Processing file with extension - ", extension)

        if extension == ".xlsx":
            print("Opening Excel file")
            return pd.read_excel(BytesIO(file))
        elif extension == ".csv":
            return pd.read_csv(BytesIO(file))
        elif extension == ".txt":
            print("Reading text file")
            return file.decode("utf-8")
        elif extension == ".docx":
            doc = Document(BytesIO(file))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        elif extension == ".pdf":
            return self.read_pdf(BytesIO(file), self.document_analysis_client)
        else:
            print("Converting file to Markdown")
            with tempfile.NamedTemporaryFile(delete=True, suffix=extension) as tmpfile:
                tmpfile.write(file)
                tmpfile.seek(0)
                return pypandoc.convert_file(tmpfile.name, "markdown")

    def read_and_validate_file(self, encoded_file: str, extension: str) -> Any:
        """
        The function reads and validates a base64 encoded file, then processes it based on the specified
        file extension.

        Args:
          encoded_file (str): The `encoded_file` parameter is a string that represents the file content
        encoded in base64 format. This function reads and validates the file content by decoding it from
        base64 and then passing it to the `read_file` method for further processing based on the specified
        file extension. If any errors occur during
          extension (str): Extension refers to the file format or type of the file being processed. It is
        typically represented by a file extension such as ".txt", ".pdf", ".jpg", etc. In the context of the
        provided code snippet, the extension parameter is used to specify the file format of the decoded
        file before further

        Returns:
          The `read_and_validate_file` method returns the output of the `read_file` method if it is not
        None. If the `read_file` method returns None, a HTTPException with status code 422 and detail
        "Failed to process the file" is raised. If any other exception occurs during the process, a
        HTTPException with status code 500 and the exception message is raised.
        """
        try:
            file_bytes = base64.b64decode(encoded_file)
            output = self.process_file_bytes(file_bytes, extension)
            if output is None:
                raise HTTPException(
                    status_code=422, detail="Failed to process the file"
                )
            return output
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e)) from e


class BytesToDocx(FastAPIUploadManager):
    def __init__(self, background_tasks):
        super().__init__(background_tasks)

    def process_file_bytes(self, file: bytes, extension=".docx") -> Document:
        """
        The function `process_file_bytes` processes a file in bytes format, specifically for a .docx
        extension, and returns a Document object while also handling background tasks.

        Args:
          file (bytes): The `file` parameter in the `process_file_bytes` method is expected to be a bytes
        object representing the content of a file. This method processes the file bytes, specifically for a
        `.docx` file. If the provided extension is not `.docx`, a `TypeError` is raised.
          extension: The `extension` parameter in the `process_file_bytes` method is used to specify the
        file extension that the method expects the input file to have. In this case, the default extension
        is set to ".docx". If the provided extension does not match ".docx", a TypeError is raised. Defaults
        to .docx

        Returns:
          The function `process_file_bytes` returns the content of a document (cv_in_docx) after processing
        a file given as bytes input.
        """
        if extension != ".docx":
            raise TypeError
        cv_in_docx_filepath, cv_in_docx = ingest_docx_bytes(file)
        self.background_tasks.add_task(os.unlink, cv_in_docx_filepath)
        return cv_in_docx
