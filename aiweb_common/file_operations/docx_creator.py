import io
import os
from docx import Document
from docx.shared import Inches
from fastapi import BackgroundTasks
from collections.abc import Mapping

import pandas as pd

from .file_handling import file_to_base64, markdown_to_docx_temporary_file


class DocxCreator:
    """
    Parent class that holds the common methods for creating DOCX reports,
    including the results summary and figures.
    """

    def __init__(self, summary = None, results=None, figures=None):
        self.results = results
        self.figures = figures
        self.summary = summary

 
    def _add_table(self, doc: Document, table_data: pd.DataFrame, heading: str, style: str = "LightShading-Accent1"):
        """
        Insert a table into the DOCX document.
        """
        doc.add_heading(heading, level=2)
        table = doc.add_table(rows=1, cols=len(table_data.columns))
        table.style = style

        # Header row.
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(table_data.columns):
            hdr_cells[i].text = col_name

        # Populate the table with data.
        for _, row in table_data.iterrows():
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row):
                row_cells[i].text = str(cell_value)

        doc.add_paragraph()

    def _add_figures(self, doc: Document):
        # Add figures / confusion matrices.
        if self.figures:
            doc.add_heading("Figures", level=1)
            for method, cm_fig in self.figures.items():
                doc.add_heading(method, level=2)
                buf = io.BytesIO()
                cm_fig.savefig(buf, format="png", bbox_inches="tight")
                buf.seek(0)
                doc.add_picture(buf, width=Inches(5))
                doc.add_paragraph()

    def _add_results_to_docx(
        self,
        doc: Document,
        results: Mapping[str, Mapping[str, object]],
        *,
        order: tuple[str, ...] | None = None,
        pretty_names: Mapping[str, str] | None = None,
    ) -> None:
        """
        Insert *all* results into the DOCX, no hard-coded section names required.

        Parameters
        ----------
        doc : python-docx Document
            The document to write into.
        results : dict[str, dict[str, Any]]
            Outer keys = method names.  
            Inner keys = arbitrary section names whose values can be rendered
            by ``self._add_table``.
        order : tuple[str] | None, default None
            If given, section keys in this tuple are shown first and in
            exactly this order; any remaining keys follow in their natural order.
        pretty_names : dict[str, str] | None, default None
            Optional mapping that converts raw section keys into nicer
            titles.  Anything not listed falls back to `title_case(key)`.
        """
        pretty_names = pretty_names or {}
        order = order or ()

        doc.add_heading("Results", level=1)
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(self.summary)
        doc.add_paragraph()
        for method, sections in results.items():
            doc.add_heading(str(method), level=2)

            # 1️⃣ keys the caller explicitly asked for, in order …
            for key in order:
                if key in sections:
                    self._add_table(doc, sections[key], pretty_names.get(key, self._title_case(key)))

            # 2️⃣ any leftover keys (excluding those already shown) …
            for key, data in sections.items():
                if key in order:
                    continue
                self._add_table(doc, data, pretty_names.get(key, self._title_case(key)))

            doc.add_paragraph()  # blank line between methods

    @staticmethod
    def _title_case(text: str) -> str:
        """Generic fallback like 'bootstrap_confidence_intervals' → 'Bootstrap Confidence Intervals'."""
        return text.replace("_", " ").title()

    def create_docx_report(self) -> Document:
        """
        Create a DOCX report that includes a heading, the results section,
        and any associated figures.
        """
        doc = Document()
        doc.add_heading("Model Evaluation Report", 0)

        # Add the method comparisons / metrics overview.
        if self.results:
            self._add_results_to_docx(doc, self.results)



        return doc

class StreamlitDocxCreator(DocxCreator):
    """
    A specialized DocxCreator for Streamlit usage. 
    It inherits the _add_results_to_docx and create_docx_report methods 
    directly from the parent class now, so there's no need to redefine them.
    """

    def __init__(self, summary, results, figures):
        super().__init__(summary=summary, results=results, figures=figures)
        # Any additional Streamlit-specific logic can go here.
        # The DOCX generation is handled by the parent.


class FastAPIDocxCreator(DocxCreator):
    """
    A specialized DocxCreator for FastAPI usage. 
    Inherits the common DOCX-report functionality and adds a method 
    to handle markdown-to-docx conversion.
    """

    def __init__(self, background_tasks: BackgroundTasks, results=None, figures=None):
        super().__init__(results=results, figures=figures)
        self.background_tasks = background_tasks

    def convert_markdown_to_docx_bytes(self, generated_response, template_filepath=None) -> str:
        """
        Convert a markdown-based response to a temporary DOCX, then
        base64-encode it for returning via FastAPI.
        """
        print("Converting markdown to temporary file")
        temp_file_path = markdown_to_docx_temporary_file(
            generated_response, template_location=template_filepath
        )
        encoded_file = file_to_base64(temp_file_path)  # Convert file to base64
        print("File encoded")
        self.background_tasks.add_task(os.unlink, temp_file_path)
        return encoded_file